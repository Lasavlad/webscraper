from bs4 import BeautifulSoup as BS
import requests
from docx import Document
import time

#allow user to filter out skills
print("Input a skill you are not familiar with:")
unfamiliar_skill = input('>')
if ' ' in unfamiliar_skill or ', ' in unfamiliar_skill:
    unfamiliar_skill_list = unfamiliar_skill.split(',')
else:
    unfamiliar_skill_list = [unfamiliar_skill]
print(f"Filtering out {unfamiliar_skill}....")

def find_jobs():
    html_content = requests.get('https://www.timesjobs.com/candidate/job-search.html?searchType=personalizedSearch&from=submit&txtKeywords=python&txtLocation=').text

    document = Document()
    soup = BS(html_content, 'lxml')
    jobs = soup.find_all('li', class_ = 'clearfix job-bx wht-shd-bx')
    number = 0
    for index, job in enumerate(jobs):
        published_date = job.find('span', class_ = 'sim-posted')
        if 'jobs-status covid-icon clearfix' in published_date:
            split = published_date.text.split()[3:]
            date = ' '.join(split)
        else:
            date = published_date.text
        if 'few' in date:
            company_name = job.find('h3', class_='joblist-comp-name').text.replace(' ','')
            if "(MoreJobs)" in company_name:
                split_name = company_name.split('\n')
                list_items = [item.strip() for item in split_name]
                refined_company_name = list_items[1]
            else:
                refined_company_name = company_name
            skills = job.find('span', class_ = 'srp-skills').text.replace(' ', '').split(",")
            skills_list = [item.strip().lower() for item in skills]
            more_info = job.header.h2.a['href']
            if all(elem not in skills_list for elem in unfamiliar_skill_list):
                with open(f'posts/{index}.txt', 'w') as f:

                    f.write(f"Company Name: {refined_company_name.strip()} \n")
                    f.write(f"Skills: {', '.join(skills_list)} \n" )
                    f.write(f'More Info: {more_info}')
                document.add_paragraph(f"Company Name: {refined_company_name.strip()}")
                document.add_paragraph(f"Skills: {', '.join(skills_list)}")
                document.add_paragraph(f'More Info: {more_info}')
                print(f'File saved: {index}')
    document.save("output.docx")
        

if __name__ == '__main__':
    while True:
        find_jobs()
        time_wait = 10
        print(f'Waiting {time_wait} minutes... ')
        time.sleep(time_wait * 60)